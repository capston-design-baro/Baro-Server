"""고소장 DOCX 파일 생성 서비스"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from typing import Dict, Any
from datetime import datetime


class ComplaintDocxService:
    """고소장 DOCX 파일 생성 서비스"""

    # 죄목 매핑
    CRIME_TYPE_MAP = {
        "fraud": "사기죄",
        "insult": "모욕죄"
    }

    def __init__(self):
        pass

    def create_complaint_docx(
        self,
        complainant_info: Dict[str, str],
        accused_info: Dict[str, str],
        crime_type: str,
        criminal_facts: str,
        accusation_reason: str,
        duplicate_complaint: bool = False,
        related_criminal_case: bool = False,
        related_civil_case: bool = False
    ) -> BytesIO:
        """
        고소장 DOCX 파일 생성

        Args:
            complainant_info: 고소인 정보 {"name": "", "address": "", "phone": ""}
            accused_info: 피고소인 정보 {"name": "", "address": "", "phone": ""}
            crime_type: 죄목 ("fraud" or "insult")
            criminal_facts: 범죄사실
            accusation_reason: 고소이유

        Returns:
            BytesIO: DOCX 파일 바이너리 스트림
        """
        doc = Document()

        # 문서 기본 설정
        self._set_document_margins(doc)

        # 페이지 번호 추가
        self._add_page_numbers(doc)

        # 1. 제목
        self._add_title(doc, crime_type)

        # 2. 고소인 정보
        self._add_complainant_info(doc, complainant_info)

        # 3. 피고소인 정보
        self._add_accused_info(doc, accused_info)

        # 4. 고소취지
        self._add_complaint_purpose(doc, crime_type)

        # 5. 범죄사실
        self._add_criminal_facts(doc, criminal_facts)

        # 6. 고소이유
        self._add_accusation_reason(doc, accusation_reason)

        # 7. 증거자료
        self._add_evidence_section(doc)

        # 8. 관련사건의 수사 및 재판 여부
        self._add_related_cases_section(doc, duplicate_complaint, related_criminal_case, related_civil_case)

        # 9. 맺음말
        self._add_closing(doc, complainant_info.get("name", "고소인"))

        # BytesIO로 저장
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return file_stream

    def _set_document_margins(self, doc: Document):
        """문서 여백 설정"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

    def _set_batang_font(self, run, bold=False, size=15):
        """바탕체 폰트 설정 헬퍼"""
        run.font.name = "바탕"
        run.font.bold = bold
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '바탕')

    def _add_page_numbers(self, doc: Document):
        """페이지 하단 가운데에 쪽번호 추가"""
        section = doc.sections[0]
        footer = section.footer

        # 푸터 문단 생성
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 페이지 번호 필드 추가
        run = paragraph.add_run()

        # 하이픈 추가
        run.add_text("- ")

        # 페이지 번호 필드
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)

        # 하이픈 추가
        run.add_text(" -")

        # 폰트 설정
        self._set_batang_font(run, bold=False, size=11)

    def _add_title(self, doc: Document, crime_type: str):
        """제목 추가"""
        # 제목 문단
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(20)

        # 제목 텍스트
        run = title.add_run("고  소  장")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = "바탕"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '바탕')

    def _add_complainant_info(self, doc: Document, info: Dict[str, str]):
        """고소인 정보 추가"""
        # 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run("1. 고소인")
        self._set_batang_font(run, bold=True, size=15)

        # 표 생성 (5행 x 4열)
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Table Grid'

        # 표 정렬 (가운데)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 열 너비 설정
        table.columns[0].width = Inches(1.0)  # 라벨 열 (성명, 주소 등)
        table.columns[1].width = Inches(1.5)  # 값 열1
        table.columns[2].width = Inches(1.5)  # 라벨 열2 (주민등록번호, 사무실주소)
        table.columns[3].width = Inches(2.1)  # 값 열2

        # 1행: 성명, 주민등록번호
        row = table.rows[0]
        row.cells[0].text = "성    명"
        row.cells[1].text = info.get('name', '')
        row.cells[2].text = "주민등록번호"
        row.cells[3].text = ""

        # 2행: 주소 (병합)
        row = table.rows[1]
        row.cells[0].text = "주    소"
        merged_cell = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        merged_cell.text = info.get('address', '')

        # 3행: 직업, 사무실주소
        row = table.rows[2]
        row.cells[0].text = "직    업"
        row.cells[1].text = info.get('job', '')
        row.cells[2].text = "사무실\n주소"
        row.cells[3].text = info.get('office_address', '')

        # 4행: 전화 (병합)
        row = table.rows[3]
        row.cells[0].text = "전    화"
        merged_cell = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        phone = info.get('phone', '')
        home_phone = info.get('home_phone', '')
        office_phone = info.get('office_phone', '')
        merged_cell.text = f"(휴대폰) {phone} (자택) {home_phone} (사무실) {office_phone}"

        # 5행: 이메일 (병합)
        row = table.rows[4]
        row.cells[0].text = "이메일"
        merged_cell = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        merged_cell.text = info.get('email', '')
        # 왼쪽 정렬
        for paragraph in merged_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 표 스타일 적용
        self._apply_table_style(table)

    def _add_accused_info(self, doc: Document, info: Dict[str, str]):
        """피고소인 정보 추가"""
        # 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(35)  # 고소인 표와의 간격 증가
        run = p.add_run("2. 피고소인")
        self._set_batang_font(run, bold=True, size=15)

        # 표 생성 (6행 x 4열) - 기타사항 행 추가
        table = doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'

        # 표 정렬 (가운데)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 열 너비 설정
        table.columns[0].width = Inches(1.0)  # 라벨 열 (성명, 주소 등)
        table.columns[1].width = Inches(1.5)  # 값 열1
        table.columns[2].width = Inches(1.5)  # 라벨 열2 (주민등록번호, 사무실주소)
        table.columns[3].width = Inches(2.1)  # 값 열2

        # 1행: 성명, 주민등록번호
        row = table.rows[0]
        row.cells[0].text = "성    명"
        row.cells[1].text = info.get('name', '')
        row.cells[2].text = "주민등록번호"
        row.cells[3].text = ""

        # 2행: 주소 (병합)
        row = table.rows[1]
        row.cells[0].text = "주    소"
        merged_cell = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        merged_cell.text = info.get('address', '')

        # 3행: 직업, 사무실주소
        row = table.rows[2]
        row.cells[0].text = "직    업"
        row.cells[1].text = info.get('job', '')
        row.cells[2].text = "사무실\n주소"
        row.cells[3].text = info.get('office_address', '')

        # 4행: 전화 (병합)
        row = table.rows[3]
        row.cells[0].text = "전    화"
        merged_cell = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        phone = info.get('phone', '')
        merged_cell.text = f"(휴대폰) {phone}"

        # 5행: 이메일 (병합)
        row = table.rows[4]
        row.cells[0].text = "이메일"
        merged_cell = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        merged_cell.text = info.get('email', '')
        # 왼쪽 정렬
        for paragraph in merged_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 6행: 기타사항 (병합)
        row = table.rows[5]
        row.cells[0].text = "기타사항"
        merged_cell = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        merged_cell.text = info.get('etc', '')
        # 왼쪽 정렬
        for paragraph in merged_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 표 스타일 적용
        self._apply_table_style(table)

    def _add_complaint_purpose(self, doc: Document, crime_type: str):
        """고소취지 섹션 추가"""
        crime_name = self.CRIME_TYPE_MAP.get(crime_type, crime_type)

        # 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(35)  # 피고소인 표와의 간격 증가
        run = p.add_run("3. 고소취지")
        self._set_batang_font(run, bold=True, size=15)

        # 내용
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"고소인은 피고소인을 {crime_name}로 고소하오니 처벌하여 주시기 바랍니다.")
        self._set_batang_font(run, bold=False, size=13)

        # 페이지 나누기 - 범죄사실이 다음 페이지로 넘어가도록
        doc.add_page_break()

    def _add_criminal_facts(self, doc: Document, criminal_facts: str):
        """범죄사실 섹션 추가"""
        # 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(15)
        run = p.add_run("4. 범죄사실")
        self._set_batang_font(run, bold=True, size=15)

        # 내용
        paragraphs = criminal_facts.split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_before = Pt(3)
                run = p.add_run(para_text)
                self._set_batang_font(run, bold=False, size=13)

    def _add_accusation_reason(self, doc: Document, accusation_reason: str):
        """고소이유 섹션 추가"""
        # 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(35)  # 범죄사실과의 간격 증가
        run = p.add_run("5. 고소이유")
        self._set_batang_font(run, bold=True, size=15)

        # 내용
        paragraphs = accusation_reason.split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_before = Pt(3)
                run = p.add_run(para_text)
                self._set_batang_font(run, bold=False, size=13)

    def _add_evidence_section(self, doc: Document):
        """증거자료 섹션 추가"""
        # 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(35)  # 고소이유와의 간격 증가
        run = p.add_run("6. 증거자료")
        self._set_batang_font(run, bold=True, size=15)

        # 체크박스 항목들
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        run = p.add_run("☐ 고소인은 고소인의 진술 외에 제출할 증거가 없습니다.")
        self._set_batang_font(run, bold=False, size=13)

        p = doc.add_paragraph()
        run = p.add_run("☑ 고소인은 고소인의 진술 외에 제출할 증거가 있습니다.")
        self._set_batang_font(run, bold=False, size=13)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run("☞ 증거자료의 세부내역은 별지를 작성하여 첨부합니다.")
        self._set_batang_font(run, bold=True, size=13)

    def _add_related_cases_section(self, doc: Document, duplicate_complaint: bool, related_criminal_case: bool, related_civil_case: bool):
        """관련사건의 수사 및 재판 여부 섹션 추가"""
        # 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(35)  # 증거자료와의 간격
        run = p.add_run("7. 관련사건의 수사 및 재판 여부")
        self._set_batang_font(run, bold=True, size=15)

        # 표 생성 (3행 x 2열, 복잡한 구조)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        # 표 정렬 (가운데)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 열 너비 설정 (총 6.1인치)
        table.columns[0].width = Inches(1.3)  # 첫 번째 열
        table.columns[1].width = Inches(4.8)  # 두 번째 열

        # 1행: 중복 고소 여부
        row = table.rows[0]
        cell0 = row.cells[0]
        cell0.text = "① 중복 고소 여부"
        # 첫 번째 열 가운데 정렬
        for paragraph in cell0.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 첫 번째 열 오른쪽 마진 조정 및 수직 가운데 정렬
        tc = cell0._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        tcMar = OxmlElement('w:tcMar')
        right = OxmlElement('w:right')
        right.set(qn('w:w'), '100')  # 왼쪽과 동일하게 100 twips
        right.set(qn('w:type'), 'dxa')
        tcMar.append(right)
        tcPr.append(tcMar)
        # 수직 가운데 정렬
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

        cell1 = row.cells[1]
        p = cell1.paragraphs[0]
        # duplicate_complaint 값에 따라 체크박스 표시
        check_yes = "☑" if duplicate_complaint else "☐"
        check_no = "☐" if duplicate_complaint else "☑"
        run = p.add_run(f"본 고소장과 같은 내용의 고소장을 다른 검찰청 또는 경찰서에 제출하거나 제출하였던 사실이 있습니다 {check_yes} / 없습니다 {check_no}")
        self._set_batang_font(run, bold=False, size=10)

        # 2행: 관련 형사사건 수사 유무
        row = table.rows[1]
        cell0 = row.cells[0]
        cell0.text = "② 관련 형사사건\n수사    유무"
        # 첫 번째 열 가운데 정렬
        for paragraph in cell0.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 첫 번째 열 오른쪽 마진 조정 및 수직 가운데 정렬
        tc = cell0._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        tcMar = OxmlElement('w:tcMar')
        right = OxmlElement('w:right')
        right.set(qn('w:w'), '100')
        right.set(qn('w:type'), 'dxa')
        tcMar.append(right)
        tcPr.append(tcMar)
        # 수직 가운데 정렬
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

        cell1 = row.cells[1]
        p = cell1.paragraphs[0]
        # related_criminal_case 값에 따라 체크박스 표시
        check_yes = "☑" if related_criminal_case else "☐"
        check_no = "☐" if related_criminal_case else "☑"
        run = p.add_run(f"본 고소장에 기재된 범죄사실과 관련된 사건 또는 공범에 대하여 검찰청이나 경찰서에서 수사 중에 \n있습니다 {check_yes} / 수사 중에 있지 않습니다 {check_no}")
        self._set_batang_font(run, bold=False, size=10)

        # 3행: 관련 민사소송 유무
        row = table.rows[2]
        cell0 = row.cells[0]
        cell0.text = "③ 관련 민사소송\n유    무"
        # 첫 번째 열 가운데 정렬
        for paragraph in cell0.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 첫 번째 열 오른쪽 마진 조정 및 수직 가운데 정렬
        tc = cell0._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        tcMar = OxmlElement('w:tcMar')
        right = OxmlElement('w:right')
        right.set(qn('w:w'), '100')
        right.set(qn('w:type'), 'dxa')
        tcMar.append(right)
        tcPr.append(tcMar)
        # 수직 가운데 정렬
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

        cell1 = row.cells[1]
        p = cell1.paragraphs[0]
        # related_civil_case 값에 따라 체크박스 표시
        check_yes = "☑" if related_civil_case else "☐"
        check_no = "☐" if related_civil_case else "☑"
        run = p.add_run(f"본 고소장에 기재된 범죄사실과 관련된 사건에 대하여 법원에서 민사소송중에 있습니다 {check_yes} / 민사소송 중에 있지 않습니다 {check_no}")
        self._set_batang_font(run, bold=False, size=10)

        # 표 스타일 적용
        self._apply_table_style(table, align_left=True)

    def _add_closing(self, doc: Document, complainant_name: str):
        """맺음말 추가"""
        # 안내 문구
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run("본 고소장에 기재한 내용은 고소인이 알고 있는 지식과 경험을 바탕으로\n모두 사실대로 작성하였으며, 만일 허위사실을 고소하였을 때에는 형법 제\n156조 무고죄로 처벌받을 것임을 서약합니다.")
        self._set_batang_font(run, bold=False, size=13)

        # 날짜 (숫자만 볼드)
        doc.add_paragraph()
        doc.add_paragraph()
        today = datetime.now()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 연도
        run = p.add_run(f"{today.year}")
        self._set_batang_font(run, bold=True, size=13)
        run = p.add_run("년     ")
        self._set_batang_font(run, bold=False, size=13)

        # 월
        run = p.add_run(f"{today.month}")
        self._set_batang_font(run, bold=True, size=13)
        run = p.add_run("월     ")
        self._set_batang_font(run, bold=False, size=13)

        # 일
        run = p.add_run(f"{today.day}")
        self._set_batang_font(run, bold=True, size=13)
        run = p.add_run("일")
        self._set_batang_font(run, bold=False, size=13)

        doc.add_paragraph()

        # 고소인 서명 (이름에 밑줄)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("고소인   ")
        self._set_batang_font(run, bold=False, size=13)

        # 이름을 띄어쓰기로 작성하고 볼드체 처리
        name_with_spaces = '   '.join(complainant_name)
        run = p.add_run(name_with_spaces)
        self._set_batang_font(run, bold=True, size=13)

        run = p.add_run("   (인).")
        self._set_batang_font(run, bold=False, size=13)

        doc.add_paragraph()
        doc.add_paragraph()

        # 수신처
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("○○지방검찰청 귀중")
        self._set_batang_font(run, bold=False, size=13)

    def _apply_table_style(self, table, align_left=False):
        """표 스타일 적용 (폰트, 정렬, 여백)"""
        # 폰트 및 정렬
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not align_left:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        self._set_batang_font(run, bold=False, size=11)

        # 셀 여백 설정
        self._set_table_cell_margins(table)
        # 행 높이 설정
        self._set_table_row_height(table, height_pt=20)
        # 테두리 설정 (외곽 굵게, 내부 기본)
        self._set_table_border(table)

    def _set_table_cell_margins(self, table):
        """표 셀 여백 설정"""
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblCellMar = OxmlElement('w:tblCellMar')

        for margin_name, value in [('top', 160), ('bottom', 160), ('left', 100), ('right', 100)]:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)

        tblPr.append(tblCellMar)

    def _set_table_row_height(self, table, height_pt=20):
        """표 행 높이 설정"""
        for row in table.rows:
            tr = row._element
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)

            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(height_pt * 20))
            trHeight.set(qn('w:hRule'), 'atLeast')
            trPr.append(trHeight)

    def _set_table_border(self, table, outer_border_size=12, inner_border_size=4):
        """표 테두리 설정 (외곽은 굵게, 내부는 기본 굵기)"""
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # 기존 테두리 설정 제거
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            tblPr.remove(tblBorders)

        # 새로운 테두리 설정
        tblBorders = OxmlElement('w:tblBorders')

        # 외곽 테두리 (굵게)
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(outer_border_size))  # 외곽 테두리 굵기
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # 검정색
            tblBorders.append(border)

        # 내부 테두리 (기본 굵기)
        for border_name in ['insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(inner_border_size))  # 내부 테두리 굵기
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # 검정색
            tblBorders.append(border)

        tblPr.append(tblBorders)


# 싱글톤 인스턴스
complaint_docx_service = ComplaintDocxService()
